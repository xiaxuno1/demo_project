# --------------------------------------------------
# coding=utf8
# !/usr/bin/python
# PN: test_tool
# FN: create_test_report
# Author: xiaxu
# DATA: 2023/12/4
# Description:docx文档操作封装
# ---------------------------------------------------
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt,Cm,Inches


class Singleton:
    #单例模式
    def __new__(cls, *args, **kwargs):
        if not hasattr(cls,'_instance'):
            cls._instance = super(Singleton,cls).__new__(cls)
        return cls._instance

class DOCXOP(Singleton):
    def __init__(self,path):
        """
        打开docx文件
        :param sheet_name: 数据表名
        """
        self.doc = Document(path)

    #替换关键字
    def replace_paragraphs_by_run(self,replace_keyword_dic):
        for p in self.doc.paragraphs:
            runs = p.runs
            #这里主要是runs会随意分割一段，导致关键字无法找到，因此合并组合，然后替换
            for i ,run in enumerate(runs):#记录关键字分割的位置，方便替换
                #print(run.text)
                if "#" in run.text: #关键字标记
                    count = i  #记录关键字起始位置
                    pre_partial = run.text[:run.text.find("#")]
                    tmp = run.text[run.text.find("#"):]  #找出关键字的部分，可能全，可能不全，可能超出
                    while tmp not in list(replace_keyword_dic.keys()):
                        count = count+1
                        if count+1>len(runs):  #超出段落分割数量，没有匹配到关键字，退出循环
                            #判断是否是第一次找到的关键字就超出真正关键字长度
                            unmatch_count = 0  #记录未匹配次数，判断是否满足条件
                            for key in replace_keyword_dic.keys():
                                if key in tmp:
                                    runs[i].text = runs[i].text.replace(runs[i].text,pre_partial+replace_keyword_dic[key]+
                                                                        tmp[tmp.find("key")+len(key)+1:]) #runs[i]= 前缀+关键字+后缀
                                    break
                                unmatch_count = unmatch_count+1
                            if unmatch_count>=len(replace_keyword_dic.keys()): #没有匹配到，提示
                                print("关键字没有在传入key找到，所在内容为：",tmp)
                            break
                        tmp = tmp+runs[count].text
                        runs[count].clear()  #这里有问题是如果关键字没有找到也会执行清空
                    if count+1<=len(runs):  #超过了分割数量还没有匹配上关键字，就替换
                        runs[i].text = runs[i].text.replace(runs[i].text,pre_partial+replace_keyword_dic[tmp])

    #通过段落替换，会丢失样式
    def replace_paragraphs(self,keyword,replace_keyword):
        for p in self.doc.paragraphs:
            if keyword in p.text:
                p.text = p.text.replace(keyword,replace_keyword)

    #指定内容填入表格
    def fill_in_table(self,table_index,content_list):
        # 读取指定表里面的数据
        table = self.doc.tables[table_index] #获取指定索引表格
        row = 0  #行号，从0开始
        for content in content_list[1:]: #不要表头
            table.add_row()  #增加单行
            row = row+1
            col = 0  # 列号，每新行从0开始
            for cell in content:
                table.cell(row,col).text = cell
                col = col + 1
    #向表格中添加图片
    def add_pic2excell(self,table_index,loc,pic_path,width=Cm(3.82),height=Cm(2.78)):
        table = self.doc.tables[table_index]  # 获取指定索引表格
        run = table.cell(loc[0],loc[1]).paragraphs[1] #添加获取指定段光标所在位置，附加到第一段[0]，不会清除原来的数据
        #将光标移动到指定字符
        print(run.text)
        run.add_picture(pic_path,width,height)  #在当前光标所在位置添加图片
        #设置图片的位置
        #pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  #指定段落对齐方式，居中对齐

    #生成测试日期
    def create_test_data(self,table_index,loc,pic_path,date):
        table = self.doc.tables[table_index]  # 获取指定索引表格
        table_cell = table.cell(loc[0], loc[1])
        table_cell.text = ''
        p1 = table_cell.add_paragraph("测试人：")
        p2 = table_cell.add_paragraph("确认人：")
        p3 = table_cell.add_paragraph("批准人：")
        #设置英文数字字体
        p1.style.font.name = 'Arial'
        # 设置中文字体
        p1.style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 设置字号大小
        p1.style.font.size = Pt(12)

        p1.add_run().add_picture(pic_path[0],width=Cm(2.41),height=Cm(0.95))  #光标移动到段尾，添加图片
        p1.add_run("                   日期："+date) #空格
        p2.add_run().add_picture(pic_path[1],width=Cm(2.41),height=Cm(0.95))
        p2.add_run("                   日期：" + date)
        p3.add_run().add_picture(pic_path[2],width=Cm(2.41),height=Cm(0.95))
        p3.add_run("                   日期：" + date)

    #在指定行列填写内容
    def fill_cell(self,table_index,loc,cell):
        table = self.doc.tables[table_index] #获取指定索引表格
        table.cell(loc[0],loc[1]).text = cell
        p = table.cell(loc[0],loc[1]).paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  #设置段落对齐方式 居中
        p.style.font.name = 'Arial'
        # 设置中文字体
        p.style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 设置字号大小
        p.style.font.size = Pt(12)

    #设置全体的页眉
    def page_header(self,header_text):
        #self.doc.settings.odd_and_even_pages_header_footer = True
        #self.doc.sections[0].even_page_header.paragraphs[0].text = header_text  #设置偶数页页眉
        #self.doc.sections[0].even_page_header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  #设置段落对齐方式
        self.doc.sections[0].header.paragraphs[0].text = header_text #设置页眉，不分奇偶页时，奇偶页都；区分时，为奇数页
        self.doc.sections[0].header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  #设置段落对齐方式

    # 按照行遍历表内容
    def table_contents_row(self,table_index):
        """
        根据表索引，按照行的方式遍历，返回list
        :param table_index:
        :return: 行组成的list
        """
        contents = [] #全体内容
        table = self.doc.tables[table_index] #获取指定索引表格
        for row in table.rows:
            content = []  #行内容列表
            for cell in row.cells:
                content.append(cell.text)
            contents.append(content)
        return contents

    #获取指定行的内容,行索引从0开始
    def table_content_row(self,table_index,row_index):
        """
        根据表索引，按照行的方式遍历，返回list
        :param table_index:
        :return: 行组成的list
        """
        content = [] #全体内容
        row_num = 0  #记录行号
        table = self.doc.tables[table_index] #获取指定索引表格
        for row in table.rows:
            if row_num == row_index:
                for cell in row.cells:
                    content.append(cell.text)
            row_num = row_num+1
        return content

    # 按照列遍历表内容
    def table_contents_col(self,table_index):
        """
        根据表索引，按照列的方式遍历，返回list
        :param table_index:
        :return: 行组成的list
        """
        contents = [] #全体内容
        table = self.doc.tables[table_index] #获取指定索引表格
        for col in table.cols:
            content = []  #列内容列表
            for cell in col.cells:
                content.append(cell.text)
        return contents

    #删除表格内容,会导致文档出问题
    def clear_table_content(self,table_index,loc):
        table = self.doc.tables[table_index] #获取指定索引表格
        table.add_row().cells[0].text = "添加第一行"
        table.add_row().cells[0].text = "添加第二行"
        #table.cell(loc[0],loc[1])._element.clear_content()

    # 打印docx内容
    def print_docx(self):
        for p in self.doc.paragraphs:
            for run in p.runs:
                print(run.text)

    # 保存新文件
    def save_docx(self,save_path):
        self.doc.save(save_path)


if __name__ == '__main__':
    #create_table()
    #doc = DOCXOP("..\\data\\cfg_test_report_20to10.docx")
    doc = DOCXOP("..\\data\\test_report.docx")
    for i in doc.table_contents_row(0):
        print(i)
    #doc.print_docx()
    #'\n测试人：                   日期：\n\n确认人：                   日期：\n\n批准人：                   日期：'
    #doc.replace_paragraphs_by_run({"#时间段":"2023-11-28"})
    #doc.fill_cell(4,(3,1),"测试结论要附加方式放在这")
    #doc.add_pic2excell(0,(11,1),pic_path="..\\data\\电子签名.jpg")  #添加图片
    #doc.create_test_data(0,(11,1),pic_path=["..\\data\\电子签名\\夏旭02.png","..\\data\\电子签名\\蒲国宇01.png","..\\data\\电子签名\\邓永亮.png"],date="2023/11/11")
    #doc.clear_table_content(0,(11,0))
    #doc.fill_in_table(0,)
    #删除指定页
    doc.page_header("2020型车站接入10中心CSM配置文件测试报告")
    print(doc.table_content_row(0, 6))
    doc.fill_cell(0, (0, 0), "2020型站机接入10中心配置文件测试报告")
    doc.save_docx("..\\data\\修改docx.docx")
        # print(d)